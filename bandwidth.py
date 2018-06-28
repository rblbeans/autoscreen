#!/bin/python
# -*- coding:utf-8 -*-
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
import datetime
import time
import os
from PIL import Image
import codecs
import sys
import pyocr
os.system('sh /root/autoscreen/wget_cacti.sh')
dd = str(datetime.date.today())
dpath = "/var/www/html/cacti/"+dd+'/'
dpath_export = "/var/www/html/cacti/"
png = ("68", "70", "72", "73", "76", "93", "94", "81", "82", "84", "87")
resault_list=[]
#builder = pyocr.builders.TextBuilder()
builder = pyocr.tesseract.DigitBuilder()
tools = pyocr.get_available_tools()[:]
if len(tools)==0:
    print("no ocr tool found")
    sys.exit(1)
else:
    print("Using '%s' " % (tools[0].get_name()))
for i in png:
    image = Image.open(dpath+i+".png")
    cut = i+"cut.png"
    x = 30
    y = 200
    w = 505
    h = 44
    region = image.crop((x, y, x+w, y+h))
    region.save(dpath+cut)
    cut_image = Image.open(dpath+cut)
    data = tools[0].image_to_string(cut_image, lang='eng', builder=builder)
    with codecs.open(dpath+"data.txt", 'a+',encoding='utf-8') as file_descriptor:
        builder.write_file(file_descriptor, data+'\n')
os.system('sh /root/autoscreen/cacti_value.sh')
with open(dpath+"max",'r') as fmax,open(dpath+"current", 'r') as fcur:
    for i in range(11):
        j = fcur.readline()
        k = fmax.readline()
        document = Document()
	now = u'当前值：'+j
        maxed = u'两日峰值：'+k
        resault_list.append(now+maxed)
add = document.add_paragraph
document.styles['Normal'].font.name = u'宋体'
document.styles['Heading 1'].font.name = u'宋体'
document.add_heading(u'带宽统计',level=1)
add(u'F5负载均衡设备', style='ListNumber')
add(u'192.168.253.2 周Active Connections峰值：')
add(u'192.168.253.3 周Active Connections峰值：')
add(u'万达出口带宽使用统计', style='ListNumber')
add(u'万达联通', style='ListBullet')
add(u'额定带宽：155M')
add(u'当前带宽：  M          两日峰值： M')
#add(resault_list[i])
add(u'万达电信', style='ListBullet')
add(u'额定带宽：150M')
add(u'当前带宽：  M          两日峰值： M')
#add(resault_list[i])
add(u'各IDC带宽使用统计', style='ListNumber')
add(u'鲁谷BGP', style='ListBullet')
add(u'额定带宽：500M')
add(resault_list[0])
document.add_picture(dpath+'68.png',width=Inches(6))
add(u'盐城电信', style='ListBullet')
add(u'额定带宽：2G')
add(resault_list[1])
document.add_picture(dpath+'70.png',width=Inches(6))
add(u'重庆联通', style='ListBullet')
add(u'额定带宽：4G')
add(resault_list[2])
document.add_picture(dpath+'72.png',width=Inches(6))
add(u'青岛电信', style='ListBullet')
add(u'额定带宽：2G')
add(resault_list[3])
document.add_picture(dpath+'73.png',width=Inches(6))
add(u'大兴BGP', style='ListBullet')
add(u'额定带宽：15M')
add(resault_list[4])
document.add_picture(dpath+'76.png',width=Inches(6))
add(u'中关村电信', style='ListBullet')
add(u'额定带宽：50M')
add(resault_list[5])
document.add_picture(dpath+'93.png',width=Inches(6))
add(u'中关村联通', style='ListBullet')
add(u'额定带宽：50M')
add(resault_list[6])
document.add_picture(dpath+'94.png',width=Inches(6))
add(u'泰安联通', style='ListBullet')
add(u'额定带宽：3G')
add(resault_list[7])
document.add_picture(dpath+'81.png',width=Inches(6))
add(u'7层联通', style='ListBullet')
add(u'额定带宽：300M')
add(resault_list[8])
document.add_picture(dpath+'82.png',width=Inches(6))
add(u'鲁谷电信', style='ListBullet')
add(u'额定带宽：300M')
add(resault_list[9])
document.add_picture(dpath+'84.png',width=Inches(6))
add(u'江门电信', style='ListBullet')
add(u'额定带宽：4G')
add(resault_list[10])
document.add_picture(dpath+'87.png',width=Inches(6))
document.save(dpath_export+u"带宽监控图.docx")

